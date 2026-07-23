"""
Бот для работы с текстами с поддержкой дедлайнов, уведомлений и фильтров
"""

import os
import logging
import asyncio
import re
import tempfile
from datetime import datetime, timedelta, time as datetime_time
from typing import Optional, List, Dict, Any
from psycopg2.extras import RealDictCursor  # <-- добавить, если нет

from telegram import Update, InlineKeyboardMarkup, InlineKeyboardButton, BotCommand
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    filters,
    ContextTypes,
    CallbackQueryHandler,
)
from dotenv import load_dotenv
import docx
import PyPDF2
from ai_processor import extract_assignments_from_text, normalize_assignment_from_groq

PAGE_SIZE = 5

load_dotenv()

EMPLOYEES = {
    "Ома К.К.": 1135811406,
    "Кулубеков Т.Т.": 400037831,
}

ADMINS_NAMES = [
    "Ома К.К.",
]

def get_full_name_by_id(user_id: int) -> Optional[str]:
    for name, uid in EMPLOYEES.items():
        if uid == user_id:
            return name
    return None

TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "")
ADMIN_TELEGRAM_ID = int(os.getenv("ADMIN_TELEGRAM_ID", "0") or "0")
USE_POSTGRES = os.getenv("USE_POSTGRES", "false").lower() == "true"
# USE_POSTGRES = False
DATABASE_URL = os.getenv("DATABASE_URL", "")

print(f"🔍 USE_POSTGRES = {USE_POSTGRES}")
print(f"🔍 DATABASE_URL = {DATABASE_URL}")

logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

user_assignments = {}
assignment_id_counter = 1

db_manager = None
if USE_POSTGRES and DATABASE_URL:
    try:
        from database import DatabaseManager
        db_manager = DatabaseManager(DATABASE_URL)
        print("✅ PostgreSQL подключена")
    except Exception as e:
        print(f"❌ Ошибка подключения к PostgreSQL: {e}, работаем в локальном режиме")
        USE_POSTGRES = False
else:
    print("ℹ️ Используется локальное хранилище")

def load_employees_from_db():
    global EMPLOYEES
    if USE_POSTGRES and db_manager:
        try:
            employees = db_manager.get_all_employees()
            EMPLOYEES = {emp["full_name"]: emp["telegram_id"] for emp in employees}
            print(f"✅ Загружено {len(EMPLOYEES)} сотрудников из базы")
        except Exception as e:
            print(f"❌ Ошибка загрузки сотрудников: {e}")
    else:
        print(f"ℹ️ Сотрудники загружены из локального словаря: {len(EMPLOYEES)} записей")

load_employees_from_db()

def is_user_registered(user_id: int) -> bool:
    if USE_POSTGRES and db_manager:
        employee = db_manager.get_employee_by_id(user_id)
        return employee is not None
    else:
        for name, uid in EMPLOYEES.items():
            if uid == user_id:
                return True
        return False

def is_admin(user_id: int) -> bool:
    if USE_POSTGRES and db_manager:
        employee = db_manager.get_employee_by_id(user_id)
        if employee:
            return employee.get("role") == "admin"
    else:
        for name, uid in EMPLOYEES.items():
            if uid == user_id and name in ADMINS_NAMES:
                return True
    return False

def require_registration(func):
    async def wrapper(update: Update, context: ContextTypes.DEFAULT_TYPE, *args, **kwargs):
        user_id = update.effective_user.id
        if not is_user_registered(user_id):
            await update.message.reply_text(
                "⚠️ *Для использования этой команды необходимо зарегистрироваться!*\n\n"
                "Используйте команду:\n"
                "📝 `/register Ваше ФИО`\n\n"
                "Пример: `/register Иванов И.И.`",
                parse_mode="Markdown"
            )
            return
        return await func(update, context, *args, **kwargs)
    return wrapper

# ============== ФУНКЦИИ РАБОТЫ С ФАЙЛАМИ ==============

def extract_text_from_pdf(file_path: str) -> str:
    try:
        text = ""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                text_parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    text_parts.append(" | ".join(row_cells))
        if not text_parts:
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""

def extract_text_from_pdf_with_mistral(file_path: str) -> str:
    mistral_api_key = os.getenv("MISTRAL_API_KEY")
    if not mistral_api_key:
        print("❌ Mistral API key не найден")
        return ""
    try:
        from mistralai import Mistral
        import base64
        import requests
        client = Mistral(api_key=mistral_api_key)
        with open(file_path, "rb") as pdf_file:
            pdf_bytes = pdf_file.read()
            encoded_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
        headers = {
            "Authorization": f"Bearer {mistral_api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "mistral-ocr-latest",
            "document": {
                "type": "document_url",
                "document_url": f"data:application/pdf;base64,{encoded_pdf}"
            }
        }
        response = requests.post(
            "https://api.mistral.ai/v1/ocr",
            json=payload,
            headers=headers,
            timeout=180
        )
        if response.status_code != 200:
            print(f"❌ Mistral OCR ошибка: {response.text[:200]}")
            return ""
        data = response.json()
        extracted_text = ""
        for page in data.get("pages", []):
            if page.get("markdown"):
                extracted_text += page["markdown"] + "\n\n"
        print(f"✅ Mistral извлёк {len(extracted_text)} символов")
        return extracted_text.strip()
    except ImportError:
        print("❌ Библиотека mistralai не установлена. Установи: pip install mistralai")
        return ""
    except Exception as e:
        print(f"❌ Mistral OCR исключение: {e}")
        return ""

def extract_text_from_image(file_path: str) -> str:
    mistral_api_key = os.getenv("MISTRAL_API_KEY")
    if not mistral_api_key:
        print("❌ Mistral API key не найден")
        return ""
    try:
        import base64
        import requests
        with open(file_path, "rb") as img_file:
            img_bytes = img_file.read()
            encoded_img = base64.b64encode(img_bytes).decode('utf-8')
        headers = {
            "Authorization": f"Bearer {mistral_api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "mistral-ocr-latest",
            "document": {
                "type": "image_url",
                "image_url": f"data:image/png;base64,{encoded_img}"
            }
        }
        response = requests.post(
            "https://api.mistral.ai/v1/ocr",
            json=payload,
            headers=headers,
            timeout=180
        )
        if response.status_code != 200:
            print(f"❌ Mistral OCR ошибка: {response.text[:200]}")
            return ""
        data = response.json()
        extracted_text = ""
        for page in data.get("pages", []):
            if page.get("markdown"):
                extracted_text += page["markdown"] + "\n\n"
        print(f"✅ Mistral извлёк {len(extracted_text)} символов из изображения")
        return extracted_text.strip()
    except Exception as e:
        print(f"❌ Ошибка обработки изображения: {e}")
        return ""

def extract_text(file_path: str) -> str:
    if file_path.lower().endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
        if len(text.strip()) < 100:
            print("🔄 Текста мало, пробуем Mistral OCR...")
            mistral_text = extract_text_from_pdf_with_mistral(file_path)
            if mistral_text:
                text = mistral_text
        return text
    elif file_path.lower().endswith(".docx"):
        return extract_text_from_docx(file_path)
    elif file_path.lower().endswith((".png", ".jpg", ".jpeg")):
        print("🔄 Обрабатываем изображение через Mistral OCR...")
        return extract_text_from_image(file_path)
    return ""

# ============== ФУНКЦИИ ХРАНЕНИЯ ==============

def save_assignment(user_id: int, assignment: Dict[str, Any]) -> int:
    if "owner_id" not in assignment:
        assignment["owner_id"] = user_id
    assignment.setdefault("completion_report", None)
    assignment.setdefault("completed_at", None)
    assignment.setdefault("status", "active")
    
    if USE_POSTGRES and db_manager:
        return db_manager.save_assignment(user_id, assignment)
    else:
        global assignment_id_counter
        if user_id not in user_assignments:
            user_assignments[user_id] = []
        assignment["id"] = assignment_id_counter
        assignment["created_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        assignment["status"] = "active"
        assignment["owner_id"] = user_id
        assignment["completion_report"] = None
        assignment["completed_at"] = None
        assignment_id_counter += 1
        user_assignments[user_id].append(assignment)
        return assignment["id"]

def get_user_assignments(user_id: int) -> List[Dict[str, Any]]:
    if USE_POSTGRES and db_manager:
        user_name = get_full_name_by_id(user_id)
        return db_manager.get_user_assignments(user_id, user_name)
    else:
        result = []
        user_name = get_full_name_by_id(user_id)
        for uid, assignments in user_assignments.items():
            for a in assignments:
                if a.get("owner_id") == user_id:
                    result.append(a)
                elif user_name and user_name in a.get("responsible", []):
                    result.append(a)
        return result

def delete_assignment(user_id: int, assignment_id: int) -> bool:
    if USE_POSTGRES and db_manager:
        return db_manager.delete_assignment(user_id, assignment_id)
    else:
        assignments = user_assignments.get(user_id, [])
        for i, a in enumerate(assignments):
            if a.get("id") == assignment_id:
                del assignments[i]
                return True
        return False

def update_assignment(user_id: int, assignment_id: int, field: str, value: str) -> bool:
    if USE_POSTGRES and db_manager:
        return db_manager.update_assignment(user_id, assignment_id, field, value)
    else:
        assignments = user_assignments.get(user_id, [])
        for a in assignments:
            if a.get("id") == assignment_id:
                if field == "deadline":
                    try:
                        datetime.strptime(value, "%Y-%m-%d")
                    except ValueError:
                        return False
                    a["deadline"] = value
                elif field == "responsible":
                    a["responsible"] = [v.strip() for v in value.split(",") if v.strip()]
                elif field == "short_task":
                    a["short_task"] = value
                elif field == "description":
                    a["description"] = value
                else:
                    return False
                return True
        return False

def mark_completed(user_id: int, assignment_id: int) -> bool:
    if USE_POSTGRES and db_manager:
        return db_manager.mark_completed(user_id, assignment_id)
    else:
        user_name = get_full_name_by_id(user_id)
        for uid, assignments in user_assignments.items():
            for a in assignments:
                if a.get("id") == assignment_id:
                    if a.get("owner_id") == user_id or (user_name and user_name in a.get("responsible", [])):
                        a["status"] = "completed"
                        return True
        return False

def clear_user_assignments(user_id: int) -> bool:
    if USE_POSTGRES and db_manager:
        return db_manager.clear_user_assignments(user_id)
    else:
        if user_id in user_assignments:
            del user_assignments[user_id]
            return True
        return False

# ============== ФОРМАТИРОВАНИЕ ==============

def format_assignment_short(idx: int, a: Dict[str, Any]) -> str:
    status = a.get("status", "active")
    if status == "active" and a.get("deadline"):
        try:
            d = datetime.strptime(a["deadline"], "%Y-%m-%d").date()
            if d < datetime.now().date():
                status = "overdue"
        except:
            pass
    status_emoji = {
        "active": "🟢",
        "completed": "✅",
        "overdue": "🔴",
    }.get(status, "⚪")
    deadline = format_date(a.get("deadline") or "нет")
    responsible = ", ".join(a.get("responsible", ["не указан"]))
    return f"{idx}. {status_emoji} {a.get('short_task', 'Без названия')} (id:{a.get('id')}) | Дедлайн: {deadline} | Ответств.: {responsible}"

def format_assignment_full(a: Dict[str, Any]) -> str:
    text = f"📄 {a.get('short_task', 'Без названия')}\n"
    text += f"📝 {a.get('description', '')[:200]}...\n"
    text += f"📅 Дедлайн: {format_date(a.get('deadline', 'не указан'))}\n"
    text += f"👤 Ответственные: {', '.join(a.get('responsible', ['не указан']))}\n"
    text += f"🕒 Сохранено: {a.get('created_at', 'неизвестно')}\n"
    return text

def format_date(date_str: str) -> str:
    if not date_str or date_str == "не указан" or date_str == "нет":
        return date_str
    try:
        if isinstance(date_str, str):
            if re.match(r'\d{2}-\d{2}-\d{4}', date_str):
                return date_str
            if re.match(r'\d{4}-\d{2}-\d{2}', date_str):
                d = datetime.strptime(date_str, "%Y-%m-%d")
                return d.strftime("%d-%m-%Y")
        return date_str
    except:
        return date_str
    
def parse_date(date_str: str) -> Optional[str]:
    """
    Парсит дату из строки в форматах:
    - YYYY-MM-DD
    - DD-MM-YYYY
    - DD.MM.YYYY
    Возвращает строку в формате YYYY-MM-DD или None, если не удалось распарсить.
    """
    date_str = date_str.strip()
    # Попробуем сначала YYYY-MM-DD
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        return dt.strftime("%Y-%m-%d")
    except ValueError:
        pass
    # Попробуем DD-MM-YYYY
    try:
        dt = datetime.strptime(date_str, "%d-%m-%Y")
        return dt.strftime("%Y-%m-%d")
    except ValueError:
        pass
    # Попробуем DD.MM.YYYY
    try:
        dt = datetime.strptime(date_str, "%d.%m.%Y")
        return dt.strftime("%Y-%m-%d")
    except ValueError:
        pass
    return None   

# ============== КОМАНДЫ ==============

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.effective_user.id
    is_registered = False
    if USE_POSTGRES and db_manager:
        employee = db_manager.get_employee_by_id(user_id)
        is_registered = employee is not None
    else:
        for name, uid in EMPLOYEES.items():
            if uid == user_id:
                is_registered = True
                break

    text = """
🤖 Бот для работы с текстами

🖼️ /upload — загрузить PDF, DOCX, PNG или JPG
📋 /list — список сохранённых текстов (с пагинацией)
📅 /today — тексты с дедлайном сегодня
📆 /week — тексты с дедлайном на неделе
⏰ /overdue — просроченные тексты
✅ /completed — выполненные поручения
👤 /responsible — фильтр по ответственному
📅 /deadline — фильтр по дате или диапазону
📌 /menu — кнопочное меню
🔍 /search <слово> — поиск по текстам
📊 /stats — статистика
🗑️ /delete <id> — удалить текст по ID
✏️ /edit <id> <поле> <значение> — изменить поле (deadline, responsible, short_task, description)
🗑️ /clear — удалить все свои поручения
📊 /export_csv — выгрузить поручения в CSV
✅ /done <id> — отметить поручение выполненным
📝 /register <ФИО> — зарегистрироваться для получения уведомлений
🔔 /remind <id> — напомнить о поручении

Примеры:
/edit 5 deadline 2025-12-31
/edit 5 responsible Иванов, Петров
    """
    await update.message.reply_text(text)
    if not is_registered:
        await update.message.reply_text(
            "⚠️ *Вы не зарегистрированы!*\n\n"
            "Чтобы начать пользоваться ботом и получать уведомления о поручениях, "
            "зарегистрируйтесь с помощью команды:\n\n"
            "📝 `/register Ваше ФИО`\n\n"
            "Пример: `/register Иванов И.И.`",
            parse_mode="Markdown"
        )

@require_registration
async def upload_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text("📤 Отправь PDF или DOCX файл.")

@require_registration
async def upload_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not update.message.document:
        await update.message.reply_text("Отправь файл.")
        return
    file = update.message.document
    file_name = file.file_name or "unknown"
    if not file_name.lower().endswith((".pdf", ".docx", ".png", ".jpg", ".jpeg")):
        await update.message.reply_text("⚠️ Поддерживаются: PDF, DOCX, PNG, JPG, JPEG")
        return
    await update.message.reply_text("📥 Скачиваю...")
    try:
        file_obj = await context.bot.get_file(file.file_id)
        file_path = os.path.join(os.getcwd(), file_name)
        await file_obj.download_to_drive(file_path)
        text = extract_text(file_path)
        os.remove(file_path)
        if not text or len(text.strip()) < 5:
            await update.message.reply_text("❌ Не удалось извлечь текст. Файл пуст или защищён.")
            return
        preview = text[:1000] + ("..." if len(text) > 1000 else "")

        if len(text) > 10000:
            text = text[:10000]

        print("⏳ Вызываем extract_assignments_from_text...")
        groq_data = extract_assignments_from_text(text)
        print(f"⏳ groq_data = {groq_data}")

        if groq_data and groq_data.get("assignments"):
            assignments_list = groq_data["assignments"]
            msg = f"🤖 Groq нашёл {len(assignments_list)} поручений:\n\n"
            for idx, raw in enumerate(assignments_list, 1):
                assignment = normalize_assignment_from_groq(raw)
                msg += f"{idx}. {assignment.get('short_task', 'Без названия')}\n"
                msg += f"   📅 Дедлайн: {format_date(assignment.get('deadline', 'не указан'))}\n"
                msg += f"   👤 Ответств.: {', '.join(assignment.get('responsible', ['не указан']))}\n\n"
            await update.message.reply_text(msg)

            prepared = []
            for raw in assignments_list:
                assignment = normalize_assignment_from_groq(raw)
                if groq_data.get("protocol") and groq_data["protocol"].get("number"):
                    assignment["protocol_number"] = groq_data["protocol"]["number"]
                prepared.append(assignment)

            context.user_data["pending_assignments"] = prepared
            context.user_data["pending_index"] = 0
            context.user_data["saved_assignments"] = []
            await show_review(update.message, context)

        else:
            assignment = {
                "short_task": file_name,
                "description": text,
                "deadline": None,
                "responsible": ["Не указан"],
            }
            aid = save_assignment(update.effective_user.id, assignment)
            await update.message.reply_text(
                f"ℹ️ Не удалось автоматически извлечь поручения. Сохранён текст без дедлайна.\n"
                f"✅ Сохранено! ID: {aid}\n"
                f"Извлечено {len(text)} символов.\n\n"
                f"📄 Текст:\n{preview}\n\n"
                f"Используй /edit {aid} deadline YYYY-MM-DD, чтобы добавить/изменить дедлайн."
            )

    except Exception as e:
        logger.error(f"Upload error: {e}")
        await update.message.reply_text(f"❌ Ошибка: {str(e)}")

@require_registration
async def list_assignments(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    print("✅ list_assignments ВЫЗВАНА")
    user_id = update.effective_user.id
    all_assignments = get_user_assignments(user_id)
    # Оставляем только активные (не выполненные)
    assignments = [a for a in all_assignments if a.get("status") != "completed"]
    if not assignments:
        await update.message.reply_text("📭 Нет активных поручений.")
        return
    context.user_data["list_assignments"] = assignments
    context.user_data["list_page"] = 0
    await send_list_page(update.message, context)

async def send_list_page(message, context, edit=False):
    print("✅ send_list_page ВЫЗВАНА")
    assignments = context.user_data.get("list_assignments", [])
    print(f"   assignments: {len(assignments)}")
    page = context.user_data.get("list_page", 0)
    total = len(assignments)
    if total == 0:
        if edit:
            await message.edit_text("📭 Нет текстов.")
        else:
            await message.reply_text("📭 Нет текстов.")
        return
    start = page * PAGE_SIZE
    end = min(start + PAGE_SIZE, total)
    page_assignments = assignments[start:end]
    text = f"📋 Сохранённые тексты (стр. {page+1}/{ (total-1)//PAGE_SIZE + 1 })\n\n"
    for idx, a in enumerate(page_assignments, start=1):
        text += format_assignment_short(idx, a) + "\n"

    keyboard = []
    nav_buttons = []
    if page > 0:
        nav_buttons.append(InlineKeyboardButton("⬅️", callback_data="list_prev"))
    if end < total:
        nav_buttons.append(InlineKeyboardButton("➡️", callback_data="list_next"))
    if nav_buttons:
        keyboard.append(nav_buttons)

    user_id = message.chat.id
    user_name = get_full_name_by_id(user_id)

    for a in page_assignments:
        if a.get("status") != "completed":
            if user_name and user_name in a.get("responsible", []):
                short = a.get("short_task", "Поручение")[:20]
                keyboard.append([
                    InlineKeyboardButton(f"✅ Выполнено: {short}", callback_data=f"complete_{a.get('id')}")
                ])

    keyboard.append([InlineKeyboardButton("🗑️ Удалить по ID", callback_data="list_delete_prompt")])
    reply_markup = InlineKeyboardMarkup(keyboard) if keyboard else None

    if edit:
        await message.edit_text(text, reply_markup=reply_markup)
    else:
        await message.reply_text(text, reply_markup=reply_markup)

async def list_navigation(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    action = query.data
    if action == "list_prev":
        context.user_data["list_page"] = max(0, context.user_data.get("list_page", 0) - 1)
    elif action == "list_next":
        context.user_data["list_page"] = context.user_data.get("list_page", 0) + 1
    elif action == "list_delete_prompt":
        await query.edit_message_text("Введите ID для удаления: /delete <id>")
        return
    await send_list_page(query.message, context, edit=True)

# ============== НОВАЯ ЛОГИКА ВЫПОЛНЕНИЯ С ОТЧЁТОМ ==============

async def complete_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    assignment_id = int(query.data.replace("complete_", ""))
    user_id = update.effective_user.id
    user_name = get_full_name_by_id(user_id)

    if not user_name:
        await query.edit_message_text("❌ Вы не зарегистрированы как сотрудник.")
        return

    found = False

    if USE_POSTGRES and db_manager:
        try:
            conn = db_manager.get_connection()
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute("SELECT * FROM assignments WHERE id = %s", (assignment_id,))
            row = cur.fetchone()
            cur.close()
            conn.close()
            if row:
                a = dict(row)
                if user_id != a.get("owner_id") and user_name not in a.get("responsible", []):
                    await query.edit_message_text("⛔ Вы не назначены ответственным за это поручение.")
                    return
                context.user_data["pending_complete"] = assignment_id
                context.user_data["pending_files"] = []
                await query.edit_message_text(
                    f"📝 Введите краткий отчёт о выполнении поручения:\n\n"
                    f"«{a.get('short_task', 'Без названия')}»\n\n"
                    f"Опишите, что было сделано.\n"
                    f"Можете также прикрепить фото или документы (по одному за раз).\n"
                    f"Отчёт будет отправлен после ввода текста."
                )
                found = True
        except Exception as e:
            logger.error(f"DB error in complete_callback: {e}")
            await query.edit_message_text("❌ Ошибка при обращении к базе данных.")
            return
    else:
        for uid, assignments in user_assignments.items():
            for a in assignments:
                if a.get("id") == assignment_id:
                    if user_id != uid and user_name not in a.get("responsible", []):
                        await query.edit_message_text("⛔ Вы не назначены ответственным за это поручение.")
                        return
                    context.user_data["pending_complete"] = assignment_id
                    context.user_data["pending_files"] = []
                    await query.edit_message_text(
                        f"📝 Введите краткий отчёт о выполнении поручения:\n\n"
                        f"«{a.get('short_task', 'Без названия')}»\n\n"
                        f"Опишите, что было сделано.\n"
                        f"Можете также прикрепить фото или документы (по одному за раз).\n"
                        f"Отчёт будет отправлен после ввода текста."
                    )
                    found = True
                    break
            if found:
                break

    if not found:
        await query.edit_message_text("❌ Поручение не найдено.")

async def photo_or_doc_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if context.user_data.get("pending_complete") is not None:
        await completion_report_handler(update, context)
    else:
        user_id = update.effective_user.id
        if not is_user_registered(user_id):
            await update.message.reply_text("⚠️ Зарегистрируйтесь: /register ФИО")
            return
        if not is_admin(user_id):
            await update.message.reply_text("⛔ Только начальники могут загружать документы.")
            return
        await upload_handler(update, context)

async def completion_report_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    pending_id = context.user_data.get("pending_complete")

    if pending_id is None:
        return

    if update.message.text:
        text = update.message.text.strip()
        if not text:
            await update.message.reply_text("⚠️ Отчёт не может быть пустым. Напишите текст.")
            return

        files = context.user_data.get("pending_files", [])
        context.user_data.pop("pending_complete", None)
        context.user_data.pop("pending_files", None)

        found = False
        owner_id = None
        assignment_data = None
        user_name = get_full_name_by_id(user_id)

        if not user_name:
            await update.message.reply_text("❌ Вы не зарегистрированы.")
            return

        if USE_POSTGRES and db_manager:
            try:
                conn = db_manager.get_connection()
                cur = conn.cursor(cursor_factory=RealDictCursor)
                cur.execute("SELECT * FROM assignments WHERE id = %s", (pending_id,))
                row = cur.fetchone()
                cur.close()
                conn.close()
                if row:
                    a = dict(row)
                    if user_id != a.get("owner_id") and user_name not in a.get("responsible", []):
                        await update.message.reply_text("⛔ Вы не назначены ответственным.")
                        return
                    conn = db_manager.get_connection()
                    cur = conn.cursor()
                    cur.execute("""
                        UPDATE assignments
                        SET status = 'completed', completion_report = %s, completed_at = NOW()
                        WHERE id = %s
                    """, (text, pending_id))
                    conn.commit()
                    cur.close()
                    conn.close()
                    found = True
                    owner_id = a.get("owner_id")
                    assignment_data = a
            except Exception as e:
                logger.error(f"DB error in completion_report_handler: {e}")
                await update.message.reply_text("❌ Ошибка при сохранении отчёта.")
                return
        else:
            for uid, assignments in user_assignments.items():
                for a in assignments:
                    if a.get("id") == pending_id:
                        if user_id != uid and user_name not in a.get("responsible", []):
                            await update.message.reply_text("⛔ Вы не назначены ответственным.")
                            return
                        a["status"] = "completed"
                        a["completion_report"] = text
                        a["completed_at"] = datetime.now().isoformat()
                        found = True
                        owner_id = uid
                        assignment_data = a
                        break
                if found:
                    break

        if not found:
            await update.message.reply_text("❌ Поручение не найдено.")
            return

        await update.message.reply_text("✅ Поручение отмечено как выполненное. Спасибо за отчёт!")

        if owner_id:
            try:
                await context.bot.send_message(
                    chat_id=owner_id,
                    text=f"📢 *Сотрудник {user_name} выполнил поручение:*\n\n"
                         f"📌 {assignment_data.get('short_task', 'Без названия')}\n"
                         f"📝 *Отчёт:* {text}"
                )
                for file_data in files:
                    try:
                        if file_data["type"] == "photo":
                            await context.bot.send_photo(
                                chat_id=owner_id,
                                photo=file_data["file_id"],
                                caption="🖼️ Фото к отчёту"
                            )
                        else:
                            await context.bot.send_document(
                                chat_id=owner_id,
                                document=file_data["file_id"],
                                caption=f"📎 Вложение к отчёту: {file_data.get('file_name', 'файл')}"
                            )
                    except Exception as e:
                        logger.error(f"Не удалось отправить файл владельцу {owner_id}: {e}")
                print(f"✅ Уведомление отправлено владельцу {owner_id}")
            except Exception as e:
                logger.error(f"Не удалось отправить уведомление владельцу {owner_id}: {e}")

    elif update.message.photo or update.message.document:
        if "pending_files" not in context.user_data:
            context.user_data["pending_files"] = []

        if update.message.photo:
            file_id = update.message.photo[-1].file_id
            file_type = "photo"
            file_name = "фото"
        else:
            file_id = update.message.document.file_id
            file_type = "document"
            file_name = update.message.document.file_name or "файл"

        context.user_data["pending_files"].append({
            "type": file_type,
            "file_id": file_id,
            "file_name": file_name
        })

        await update.message.reply_text(
            f"✅ Файл «{file_name}» добавлен к отчёту.\n"
            f"Можете добавить ещё файлы или напишите текст отчёта, чтобы завершить."
        )
    else:
        await update.message.reply_text("⚠️ Пожалуйста, отправьте текст или файл.")

# ============== ОСТАЛЬНЫЕ КОМАНДЫ ==============

@require_registration
async def delete_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    args = context.args
    if not args:
        await update.message.reply_text("Укажи ID: /delete 5")
        return
    try:
        aid = int(args[0])
    except ValueError:
        await update.message.reply_text("ID должен быть числом.")
        return
    if delete_assignment(user_id, aid):
        await update.message.reply_text(f"✅ Текст с ID {aid} удалён.")
    else:
        await update.message.reply_text(f"❌ Текст с ID {aid} не найден.")

@require_registration
async def edit_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    args = context.args
    if len(args) < 3:
        await update.message.reply_text(
            "Использование: /edit <id> <поле> <значение>\n"
            "Поля: deadline, responsible, short_task, description\n"
            "Пример: /edit 5 deadline 2025-12-31"
        )
        return
    try:
        aid = int(args[0])
    except ValueError:
        await update.message.reply_text("ID должен быть числом.")
        return
    field = args[1].lower()
    value = " ".join(args[2:])
    if field not in ["deadline", "responsible", "short_task", "description"]:
        await update.message.reply_text("Доступные поля: deadline, responsible, short_task, description")
        return
    if update_assignment(user_id, aid, field, value):
        await update.message.reply_text(f"✅ Поле {field} обновлено для ID {aid}.")
    else:
        await update.message.reply_text(f"❌ Не удалось обновить. Проверь ID и формат (для даты: YYYY-MM-DD).")

@require_registration
async def search_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    args = context.args
    if not args:
        await update.message.reply_text("Введите слово для поиска: /search отчёт")
        return
    keyword = " ".join(args).lower()
    assignments = get_user_assignments(user_id)
    found = []
    for a in assignments:
        if keyword in a.get("short_task", "").lower() or keyword in a.get("description", "").lower():
            found.append(a)
    if not found:
        await update.message.reply_text(f"🔍 По запросу «{keyword}» ничего не найдено.")
        return
    text = f"🔍 Результаты поиска ({len(found)}):\n\n"
    for a in found[:10]:
        text += format_assignment_short(a.get("id"), a) + "\n"
    if len(found) > 10:
        text += "\n... показаны первые 10."
    await update.message.reply_text(text)

@require_registration
async def stats_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    assignments = get_user_assignments(user_id)
    total = len(assignments)
    if total == 0:
        await update.message.reply_text("📊 Статистика\n\nНет данных.")
        return

    chars = sum(len(a.get("description", "")) for a in assignments)
    with_deadline = sum(1 for a in assignments if a.get("deadline"))
    without_deadline = total - with_deadline

    active = 0
    overdue = 0
    completed = 0
    today = datetime.now().date()
    for a in assignments:
        if a.get("status") == "completed":
            completed += 1
        elif a.get("deadline"):
            try:
                d = datetime.strptime(a["deadline"], "%Y-%m-%d").date()
                if d < today:
                    overdue += 1
                else:
                    active += 1
            except:
                active += 1
        else:
            active += 1

    text = f"📊 *Статистика*\n\n"
    text += f"Всего поручений: {total}\n"
    text += f"Активных: {active}\n"
    text += f"Просроченных: {overdue}\n"
    if completed > 0:
        text += f"Выполненных: {completed}\n"
    text += f"С дедлайном: {with_deadline}\n"
    text += f"Без дедлайна: {without_deadline}\n"
    text += f"Всего символов: {chars}\n"
    if total > 0:
        text += f"Средняя длина: {chars // total} символов."
    await update.message.reply_text(text, parse_mode="Markdown")

@require_registration
async def export_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    assignments = get_user_assignments(user_id)
    if not assignments:
        await update.message.reply_text("📭 Нет данных для экспорта.")
        return
    try:
        from openpyxl import Workbook
        import tempfile
        wb = Workbook()
        ws = wb.active
        ws.title = "Поручения"
        ws.append(["ID", "Краткое описание", "Дедлайн", "Ответственные", "Описание", "Создано"])
        for a in assignments:
            ws.append([
                a.get("id"),
                a.get("short_task"),
                a.get("deadline"),
                ", ".join(a.get("responsible", [])),
                a.get("description", "")[:100],
                a.get("created_at")
            ])
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            wb.save(tmp.name)
            file_path = tmp.name
        with open(file_path, "rb") as f:
            await update.message.reply_document(
                document=f,
                filename=f"поручения_{user_id}.xlsx"
            )
        await asyncio.sleep(0.5)
        try:
            os.remove(file_path)
        except:
            pass
    except ImportError:
        await update.message.reply_text("❌ Библиотека openpyxl не установлена. Установи: pip install openpyxl")
    except Exception as e:
        await update.message.reply_text(f"❌ Ошибка экспорта: {str(e)}")

@require_registration
async def export_csv_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    assignments = get_user_assignments(user_id)
    if not assignments:
        await update.message.reply_text("📭 Нет данных для экспорта.")
        return
    try:
        import csv
        import tempfile
        with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False, encoding="utf-8-sig") as tmp:
            writer = csv.writer(tmp, delimiter=";")
            writer.writerow(["ID", "Краткое описание", "Дедлайн", "Ответственные", "Описание", "Создано", "Статус"])
            for a in assignments:
                writer.writerow([
                    a.get("id"),
                    a.get("short_task"),
                    a.get("deadline"),
                    ", ".join(a.get("responsible", [])),
                    a.get("description", "")[:100],
                    a.get("created_at"),
                    a.get("status", "active")
                ])
            file_path = tmp.name
        with open(file_path, "rb") as f:
            await update.message.reply_document(
                document=f,
                filename=f"поручения_{user_id}.csv"
            )
        await asyncio.sleep(0.5)
        try:
            os.remove(file_path)
        except:
            pass
    except Exception as e:
        await update.message.reply_text(f"❌ Ошибка экспорта в CSV: {str(e)}")

@require_registration
async def clear_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if clear_user_assignments(user_id):
        await update.message.reply_text("✅ Все ваши поручения удалены.")
    else:
        await update.message.reply_text("📭 У вас нет сохранённых поручений.")

@require_registration
async def done_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    args = context.args
    if not args:
        await update.message.reply_text("Укажи ID поручения: /done 5")
        return
    try:
        assignment_id = int(args[0])
    except ValueError:
        await update.message.reply_text("ID должен быть числом.")
        return

    user_name = get_full_name_by_id(user_id)
    if not user_name:
        await update.message.reply_text("❌ Вы не зарегистрированы как сотрудник.")
        return

    found = False
    if USE_POSTGRES and db_manager:
        try:
            conn = db_manager.get_connection()
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute("SELECT * FROM assignments WHERE id = %s", (assignment_id,))
            row = cur.fetchone()
            cur.close()
            conn.close()
            if row:
                a = dict(row)
                if a.get("owner_id") == user_id or user_name in a.get("responsible", []):
                    if a.get("status") == "completed":
                        await update.message.reply_text("ℹ️ Поручение уже выполнено.")
                        return
                    conn = db_manager.get_connection()
                    cur = conn.cursor()
                    cur.execute("UPDATE assignments SET status = 'completed', completed_at = NOW() WHERE id = %s", (assignment_id,))
                    conn.commit()
                    cur.close()
                    conn.close()
                    found = True
                    await update.message.reply_text(f"✅ Поручение «{a.get('short_task', '')}» отмечено как выполненное.")
        except Exception as e:
            logger.error(f"DB error in done_command: {e}")
            await update.message.reply_text("❌ Ошибка при обращении к базе данных.")
            return
    else:
        for uid, assignments in user_assignments.items():
            for a in assignments:
                if a.get("id") == assignment_id:
                    if a.get("owner_id") == user_id or (user_name and user_name in a.get("responsible", [])):
                        if a.get("status") == "completed":
                            await update.message.reply_text("ℹ️ Поручение уже выполнено.")
                            return
                        a["status"] = "completed"
                        await update.message.reply_text(f"✅ Поручение «{a.get('short_task', '')}» отмечено как выполненное.")
                        found = True
                        break
            if found:
                break

    if not found:
        await update.message.reply_text("❌ Поручение не найдено или у вас нет прав.")

@require_registration
async def remind_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    args = context.args
    if not args:
        await update.message.reply_text("Укажи ID поручения: /remind 5")
        return
    try:
        assignment_id = int(args[0])
    except ValueError:
        await update.message.reply_text("ID должен быть числом.")
        return

    assignments = get_user_assignments(user_id)
    for a in assignments:
        if a.get("id") == assignment_id:
            text = f"🔔 *Напоминание о поручении*\n\n"
            text += f"📌 *{a.get('short_task', 'Без названия')}*\n"
            text += f"📝 {a.get('description', 'Описание отсутствует')}\n"
            if a.get("deadline"):
                text += f"📅 Дедлайн: {format_date(a.get('deadline'))}\n"
                try:
                    days_left = (datetime.strptime(a["deadline"], "%Y-%m-%d").date() - datetime.now().date()).days
                    if days_left < 0:
                        text += f"⚠️ *ПРОСРОЧЕНО!* (на {-days_left} дн.)\n"
                    elif days_left == 0:
                        text += f"🔥 *СЕГОДНЯ!*\n"
                    elif days_left == 1:
                        text += f"⏰ *ЗАВТРА!*\n"
                    else:
                        text += f"⏳ Осталось дней: {days_left}\n"
                except:
                    pass
            else:
                text += f"📅 Дедлайн не указан\n"
            text += f"👤 Ответственные: {', '.join(a.get('responsible', ['не указан']))}\n"
            await update.message.reply_text(text, parse_mode="Markdown")
            await update.message.reply_text("✅ Напоминание отправлено.")
            return
    await update.message.reply_text(f"❌ Поручение с ID {assignment_id} не найдено.")

@require_registration
async def completed_assignments(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    assignments = get_user_assignments(user_id)
    completed = [a for a in assignments if a.get("status") == "completed"]
    if not completed:
        await update.message.reply_text("📭 Нет выполненных поручений.")
        return
    text = "✅ *Выполненные поручения*\n\n"
    for idx, a in enumerate(completed, 1):
        text += format_assignment_short(idx, a) + "\n"
    await update.message.reply_text(text, parse_mode="Markdown")

@require_registration
async def today_assignments(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await filter_by_deadline(update, context, days=0)

@require_registration
async def week_assignments(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await filter_by_deadline(update, context, days=7)

@require_registration
async def overdue_assignments(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await filter_by_deadline(update, context, days=None, overdue=True)

async def filter_by_deadline(update: Update, context: ContextTypes.DEFAULT_TYPE, days: Optional[int] = None, overdue: bool = False):
    user_id = update.effective_user.id
    assignments = get_user_assignments(user_id)
    today = datetime.now().date()
    if overdue:
        filtered = [a for a in assignments if a.get("deadline") and datetime.strptime(a["deadline"], "%Y-%m-%d").date() < today]
    elif days == 0:
        filtered = [a for a in assignments if a.get("deadline") == today.isoformat()]
    elif days is not None:
        week_later = today + timedelta(days=days)
        filtered = [a for a in assignments if a.get("deadline") and today <= datetime.strptime(a["deadline"], "%Y-%m-%d").date() <= week_later]
    else:
        filtered = []
    if not filtered:
        await update.message.reply_text("📭 Нет текстов, подходящих под этот фильтр.")
        return
    text = f"📋 Результаты фильтрации\n\n"
    for idx, a in enumerate(filtered, 1):
        text += format_assignment_short(idx, a) + "\n"
    await update.message.reply_text(text)

@require_registration
async def responsible_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    assignments = get_user_assignments(user_id)
    names = set()
    for a in assignments:
        for r in a.get("responsible", []):
            if r and r != "Не указан":
                names.add(r)
    if not names:
        await update.message.reply_text("📭 Нет активных поручений с ответственными.")
        return
    keyboard = []
    row = []
    for name in sorted(names):
        row.append(InlineKeyboardButton(name, callback_data=f"resp_{name}"))
        if len(row) == 2:
            keyboard.append(row)
            row = []
    if row:
        keyboard.append(row)
    keyboard.append([InlineKeyboardButton("↩️ Назад", callback_data="menu_list")])
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("👤 *Выбери ответственного:*", reply_markup=reply_markup, parse_mode="Markdown")

async def responsible_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    name = query.data.replace("resp_", "")
    user_id = update.effective_user.id
    assignments = get_user_assignments(user_id)
    filtered = [a for a in assignments if any(r == name for r in a.get("responsible", []))]
    if not filtered:
        await query.edit_message_text(f"📭 Нет поручений у {name}.")
        return
    text = f"👤 *Поручения по ответственному: {name}*\n\n"
    for idx, a in enumerate(filtered, 1):
        text += format_assignment_short(idx, a) + "\n"
    await query.edit_message_text(text, parse_mode="Markdown")

@require_registration
async def deadline_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    args = context.args
    if not args:
        await update.message.reply_text(
            "Укажи дату в формате YYYY-MM-DD или DD-MM-YYYY (или диапазон):\n"
            "Пример: /deadline 2026-07-15\n"
            "или /deadline 15-07-2026\n"
            "или /deadline 2026-07-10 2026-07-20"
        )
        return

    # Парсим все аргументы
    parsed_dates = []
    for arg in args:
        parsed = parse_date(arg)
        if parsed is None:
            await update.message.reply_text(f"❌ Неверный формат даты: {arg}. Используй YYYY-MM-DD или DD-MM-YYYY.")
            return
        parsed_dates.append(parsed)

    try:
        if len(parsed_dates) == 1:
            date_yyyy = parsed_dates[0]
            display_date = datetime.strptime(date_yyyy, "%Y-%m-%d").strftime("%d-%m-%Y")
            assignments = get_user_assignments(user_id)
            filtered = [a for a in assignments if a.get("deadline") == date_yyyy]
            title = f"📅 Поручения с дедлайном {display_date}"
        elif len(parsed_dates) == 2:
            date_from = parsed_dates[0]
            date_to = parsed_dates[1]
            if date_from > date_to:
                date_from, date_to = date_to, date_from
            display_from = datetime.strptime(date_from, "%Y-%m-%d").strftime("%d-%m-%Y")
            display_to = datetime.strptime(date_to, "%Y-%m-%d").strftime("%d-%m-%Y")
            assignments = get_user_assignments(user_id)
            filtered = [a for a in assignments if a.get("deadline") and date_from <= a.get("deadline") <= date_to]
            title = f"📅 Поручения с дедлайном с {display_from} по {display_to}"
        else:
            await update.message.reply_text("Слишком много аргументов. Укажи одну дату или диапазон из двух дат.")
            return
    except Exception as e:
        logger.error(f"Deadline command error: {e}")
        await update.message.reply_text("❌ Ошибка обработки дат.")
        return

    if not filtered:
        await update.message.reply_text("📭 Нет поручений в указанном диапазоне.")
        return

    text = f"{title}\n\n"
    for idx, a in enumerate(filtered, 1):
        text += format_assignment_short(idx, a) + "\n"
    await update.message.reply_text(text)

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = """🤖 Справка по командам бота

🖼️ /upload — загрузить PDF, DOCX, PNG или JPG
📋 /list — список активных поручений
🔥 /today — поручения на сегодня
⏳ /week — поручения на неделю
🔴 /overdue — просроченные
✅ /completed — выполненные
👤 /responsible — фильтр по ответственному
📅 /deadline — фильтр по дате (например, /deadline 2026-07-15)
📌 /menu — кнопочное меню
✏️ /edit <id> <поле> <значение> — редактировать
🗑️ /delete <id> — удалить
🔍 /search <слово> — поиск
📊 /stats — статистика
📊 /export — выгрузка в Excel
📊 /export_csv — выгрузка в CSV
🗑️ /clear — удалить все
✅ /done <id> — отметить выполненным
🔔 /remind <id> — напомнить
📝 /register <ФИО> — зарегистрироваться для получения уведомлений
❓ /help — эта справка"""
    await update.message.reply_text(text)

@require_registration
async def menu_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("📤 Загрузить документ", callback_data="menu_upload")],
        [InlineKeyboardButton("📋 Активные", callback_data="menu_list")],
        [InlineKeyboardButton("🔥 Сегодня", callback_data="menu_today")],
        [InlineKeyboardButton("⏳ Неделя", callback_data="menu_week")],
        [InlineKeyboardButton("🔴 Просроченные", callback_data="menu_overdue")],
        [InlineKeyboardButton("✅ Выполненные", callback_data="menu_completed")],
        [InlineKeyboardButton("👤 По ответственному", callback_data="menu_responsible")],
        [InlineKeyboardButton("📅 По дедлайну", callback_data="menu_deadline")],
        [InlineKeyboardButton("📊 Экспорт CSV", callback_data="menu_export_csv")],
        [InlineKeyboardButton("❓ Помощь", callback_data="menu_help")],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("📌 *Выбери действие:*", reply_markup=reply_markup, parse_mode="Markdown")

async def menu_callback_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    action = query.data

    class FakeUpdate:
        effective_user = update.effective_user
        effective_message = update.effective_message
        message = update.effective_message

    fake_update = FakeUpdate()

    if action == "menu_upload":
        await upload_command(fake_update, context)
    elif action == "menu_list":
        await list_assignments(fake_update, context)
    elif action == "menu_today":
        await today_assignments(fake_update, context)
    elif action == "menu_week":
        await week_assignments(fake_update, context)
    elif action == "menu_overdue":
        await overdue_assignments(fake_update, context)
    elif action == "menu_completed":
        await completed_assignments(fake_update, context)
    elif action == "menu_responsible":
        await responsible_command(fake_update, context)
    elif action == "menu_deadline":
        await query.edit_message_text("Введите дату: /deadline YYYY-MM-DD")
        return
    elif action == "menu_export_csv":
        await export_csv_command(fake_update, context)
    elif action == "menu_help":
        await help_command(fake_update, context)

    await query.edit_message_reply_markup(reply_markup=None)

# ============== ДЕДУПЛИКАЦИЯ ==============

def is_duplicate(new: Dict[str, Any], existing: Dict[str, Any]) -> bool:
    new_text = new.get("short_task", "").lower().replace(" ", "")
    existing_text = existing.get("short_task", "").lower().replace(" ", "")
    if new_text == existing_text:
        return True
    new_deadline = new.get("deadline")
    existing_deadline = existing.get("deadline")
    if new_deadline and existing_deadline and new_deadline == existing_deadline:
        new_resp = set(new.get("responsible", []))
        existing_resp = set(existing.get("responsible", []))
        if new_resp and existing_resp and new_resp.intersection(existing_resp):
            return True
    return False

def find_duplicates(new_assignment: Dict[str, Any], existing_assignments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    duplicates = []
    for existing in existing_assignments:
        if is_duplicate(new_assignment, existing):
            duplicates.append(existing)
    return duplicates

async def show_duplicate_review(message_or_query, context, edit=False):
    pending_duplicates = context.user_data.get("pending_duplicates", [])
    index = context.user_data.get("pending_duplicates_index", 0)
    if not pending_duplicates:
        await message_or_query.reply_text("Все дубли обработаны.")
        return

    new_assignment, duplicates = pending_duplicates[index]
    total = len(pending_duplicates)

    text = f"🔍 *Найден возможный дубль ({index+1}/{total})*\n\n"
    text += f"📌 *Новое поручение:*\n"
    text += f"   {new_assignment.get('short_task', 'Без названия')}\n"
    text += f"   Дедлайн: {new_assignment.get('deadline', 'не указан')}\n"
    text += f"   Ответственные: {', '.join(new_assignment.get('responsible', ['не указан']))}\n"
    if new_assignment.get("protocol_number"):
        text += f"   Протокол №{new_assignment.get('protocol_number')}\n"
    text += f"\n📋 *Существующее поручение:*\n"
    existing = duplicates[0]
    text += f"   {existing.get('short_task', 'Без названия')}\n"
    text += f"   Дедлайн: {existing.get('deadline', 'не указан')}\n"
    text += f"   Ответственные: {', '.join(existing.get('responsible', ['не указан']))}\n"
    if existing.get("protocol_number"):
        text += f"   Протокол №{existing.get('protocol_number')}\n"

    keyboard = [
        [
            InlineKeyboardButton("💾 Сохранить как новое", callback_data=f"dup_save_{index}"),
            InlineKeyboardButton("🔄 Обновить существующее", callback_data=f"dup_update_{index}"),
        ],
        [InlineKeyboardButton("⏭️ Пропустить", callback_data=f"dup_skip_{index}")],
    ]
    if total > 1:
        keyboard.append([InlineKeyboardButton("➡️ Следующий дубль", callback_data="dup_next")])

    reply_markup = InlineKeyboardMarkup(keyboard)
    if edit:
        await message_or_query.edit_message_text(text, reply_markup=reply_markup, parse_mode="Markdown")
    else:
        await message_or_query.reply_text(text, reply_markup=reply_markup, parse_mode="Markdown")

async def duplicate_callback_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    action = query.data
    user_id = update.effective_user.id

    pending_duplicates = context.user_data.get("pending_duplicates", [])
    index = context.user_data.get("pending_duplicates_index", 0)

    if action == "dup_next":
        context.user_data["pending_duplicates_index"] = min(index + 1, len(pending_duplicates) - 1)
        await show_duplicate_review(query, context, edit=True)
        return

    if index >= len(pending_duplicates):
        await query.edit_message_text("Все дубли обработаны.")
        return

    new_assignment, duplicates = pending_duplicates[index]
    existing = duplicates[0]

    if action.startswith("dup_save_"):
        save_assignment(user_id, new_assignment)
        await query.edit_message_text("✅ Поручение сохранено как новое.")
    elif action.startswith("dup_update_"):
        user_assignments_list = user_assignments.get(user_id, [])
        for a in user_assignments_list:
            if a.get("id") == existing.get("id"):
                a["short_task"] = new_assignment.get("short_task", a["short_task"])
                a["description"] = new_assignment.get("description", a["description"])
                a["deadline"] = new_assignment.get("deadline", a["deadline"])
                a["responsible"] = new_assignment.get("responsible", a["responsible"])
                if new_assignment.get("protocol_number"):
                    a["protocol_number"] = new_assignment.get("protocol_number")
                await query.edit_message_text("✅ Поручение обновлено.")
                break
        else:
            await query.edit_message_text("❌ Не удалось найти поручение для обновления.")
    elif action.startswith("dup_skip_"):
        await query.edit_message_text("⏭️ Поручение пропущено.")

    pending_duplicates.pop(index)
    context.user_data["pending_duplicates"] = pending_duplicates
    context.user_data["pending_duplicates_index"] = min(index, len(pending_duplicates) - 1) if pending_duplicates else 0

    if pending_duplicates:
        await show_duplicate_review(query, context, edit=True)
    else:
        await query.edit_message_text("✅ Все дубли обработаны.")

# ============== ПРОСМОТР ПЕРЕД СОХРАНЕНИЕМ ==============

async def show_review(message_or_query, context, edit=False):
    pending = context.user_data.get("pending_assignments", [])
    index = context.user_data.get("pending_index", 0)
    if not pending:
        await message_or_query.reply_text("✅ Все поручения обработаны.")
        return

    assignment = pending[index]
    total = len(pending)

    text = f"📋 *Поручение {index+1}/{total}*\n\n"
    text += f"📌 {assignment.get('short_task', 'Без названия')}\n"
    text += f"📅 Дедлайн: {format_date(assignment.get('deadline', 'не указан'))}\n"
    text += f"👤 Ответственные: {', '.join(assignment.get('responsible', ['не указан']))}\n"
    if assignment.get("protocol_number"):
        text += f"📄 Протокол №{assignment.get('protocol_number')}\n"
    text += f"📝 Описание: {assignment.get('description', '')[:200]}..."

    keyboard = [
        [
            InlineKeyboardButton("✅ Сохранить", callback_data="review_save"),
            InlineKeyboardButton("✏️ Редактировать", callback_data="review_edit"),
        ],
        [
            InlineKeyboardButton("⏭️ Пропустить", callback_data="review_skip"),
        ],
    ]
    if total > 1:
        keyboard.append([
            InlineKeyboardButton("⬅️ Назад", callback_data="review_prev"),
            InlineKeyboardButton("➡️ Далее", callback_data="review_next"),
        ])

    reply_markup = InlineKeyboardMarkup(keyboard)

    if edit:
        await message_or_query.edit_message_text(text, reply_markup=reply_markup, parse_mode="Markdown")
    else:
        await message_or_query.reply_text(text, reply_markup=reply_markup, parse_mode="Markdown")

async def review_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    action = query.data
    user_id = update.effective_user.id

    pending = context.user_data.get("pending_assignments", [])
    index = context.user_data.get("pending_index", 0)
    saved = context.user_data.get("saved_assignments", [])

    if action == "review_prev":
        context.user_data["pending_index"] = max(0, index - 1)
        await show_review(query, context, edit=True)
        return

    if action == "review_next":
        context.user_data["pending_index"] = min(len(pending) - 1, index + 1)
        await show_review(query, context, edit=True)
        return

    if action == "review_save":
        assignment = pending[index]
        save_assignment(user_id, assignment)
        for resp in assignment.get("responsible", []):
            if resp in EMPLOYEES:
                telegram_id = EMPLOYEES[resp]
                try:
                    await context.bot.send_message(
                        chat_id=telegram_id,
                        text=f"📌 *Вы назначены ответственным за поручение:*\n\n"
                             f"📝 {assignment.get('short_task', 'Без названия')}\n"
                             f"📅 Дедлайн: {format_date(assignment.get('deadline', 'не указан'))}\n"
                             f"👤 Ответственный: {resp}"
                    )
                    print(f"✅ Уведомление отправлено {resp}")
                except Exception as e:
                    print(f"❌ Не удалось отправить уведомление {resp}: {e}")
        saved.append(assignment)
        context.user_data["saved_assignments"] = saved
        pending.pop(index)
        context.user_data["pending_assignments"] = pending
        if pending:
            context.user_data["pending_index"] = min(index, len(pending) - 1)
            await show_review(query, context, edit=True)
        else:
            await query.edit_message_text(f"✅ Сохранено поручений: {len(saved)}")
            context.user_data.pop("pending_assignments", None)
            context.user_data.pop("pending_index", None)
            context.user_data.pop("saved_assignments", None)
        return

    if action == "review_skip":
        pending.pop(index)
        context.user_data["pending_assignments"] = pending
        if pending:
            context.user_data["pending_index"] = min(index, len(pending) - 1)
            await show_review(query, context, edit=True)
        else:
            await query.edit_message_text(f"✅ Сохранено поручений: {len(saved)}")
            context.user_data.pop("pending_assignments", None)
            context.user_data.pop("pending_index", None)
            context.user_data.pop("saved_assignments", None)
        return

    if action == "review_edit":
        await query.edit_message_text(
            "✏️ Введи новое краткое описание для этого поручения.\n"
            "Или напиши 'отмена', чтобы отменить редактирование."
        )
        context.user_data["editing_index"] = index
        return

# ============== ОБЪЕДИНЁННЫЙ ОБРАБОТЧИК РЕДАКТИРОВАНИЯ ==============

async def review_edit_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    text = update.message.text.strip()
    index = context.user_data.get("editing_index")
    pending = context.user_data.get("pending_assignments", [])

    if index is None or index >= len(pending):
        return

    if text.lower() == "отмена":
        await update.message.reply_text("Редактирование отменено.")
        context.user_data.pop("editing_index", None)
        await show_review(update.message, context)
        return

    pending[index]["short_task"] = text
    context.user_data["pending_assignments"] = pending
    context.user_data.pop("editing_index", None)

    saved = context.user_data.get("saved_assignments", [])
    for assignment in pending:
        save_assignment(user_id, assignment)
        for resp in assignment.get("responsible", []):
            if resp in EMPLOYEES:
                telegram_id = EMPLOYEES[resp]
                try:
                    await context.bot.send_message(
                        chat_id=telegram_id,
                        text=f"📌 *Вы назначены ответственным за поручение:*\n\n"
                             f"📝 {assignment.get('short_task', 'Без названия')}\n"
                             f"📅 Дедлайн: {format_date(assignment.get('deadline', 'не указан'))}\n"
                             f"👤 Ответственный: {resp}"
                    )
                    print(f"✅ Уведомление отправлено {resp}")
                except Exception as e:
                    print(f"❌ Не удалось отправить уведомление {resp}: {e}")
        saved.append(assignment)

    context.user_data["saved_assignments"] = saved
    context.user_data.pop("pending_assignments", None)
    context.user_data.pop("pending_index", None)

    await update.message.reply_text(f"✅ Сохранено поручений: {len(saved)}")

# ============== ЕЖЕДНЕВНЫЕ НАПОМИНАНИЯ ==============

async def daily_reminder(context: ContextTypes.DEFAULT_TYPE):
    today = datetime.now().date()
    tomorrow = today + timedelta(days=1)

    if USE_POSTGRES and db_manager:
        all_assignments = db_manager.get_all_assignments()
    else:
        all_assignments = []
        for user_id, assignments in user_assignments.items():
            for a in assignments:
                all_assignments.append(a)

    grouped = {}
    for a in all_assignments:
        user_id = a.get("user_id")
        if user_id not in grouped:
            grouped[user_id] = []
        grouped[user_id].append(a)

    for user_id, assignments in grouped.items():
        due_today = []
        due_tomorrow = []
        overdue = []
        for a in assignments:
            if not a.get("deadline"):
                continue
            try:
                d = datetime.strptime(a["deadline"], "%Y-%m-%d").date()
            except:
                continue
            if d < today:
                overdue.append(a)
            elif d == today:
                due_today.append(a)
            elif d == tomorrow:
                due_tomorrow.append(a)

        if not (due_today or due_tomorrow or overdue):
            continue

        text = "📅 Ежедневное напоминание о дедлайнах\n\n"
        if overdue:
            text += f"🔴 Просрочено ({len(overdue)})\n"
            for a in overdue[:5]:
                text += f"• {a.get('short_task', 'Без названия')} (id:{a.get('id')})\n"
            if len(overdue) > 5:
                text += f"... и ещё {len(overdue)-5}\n"
        if due_today:
            text += f"🔥 Сегодня ({len(due_today)})\n"
            for a in due_today[:5]:
                text += f"• {a.get('short_task', 'Без названия')} (id:{a.get('id')})\n"
            if len(due_today) > 5:
                text += f"... и ещё {len(due_today)-5}\n"
        if due_tomorrow:
            text += f"⏰ Завтра ({len(due_tomorrow)})\n"
            for a in due_tomorrow[:5]:
                text += f"• {a.get('short_task', 'Без названия')} (id:{a.get('id')})\n"
            if len(due_tomorrow) > 5:
                text += f"... и ещё {len(due_tomorrow)-5}\n"

        try:
            await context.bot.send_message(chat_id=user_id, text=text)
        except Exception as e:
            logger.error(f"Не удалось отправить напоминание пользователю {user_id}: {e}")

# ============== ОБРАБОТЧИК ОШИБОК ==============

async def error_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logger.error(f"Update {update} caused error {context.error}")
    if ADMIN_TELEGRAM_ID:
        try:
            await context.bot.send_message(
                chat_id=ADMIN_TELEGRAM_ID,
                text=f"⚠️ Ошибка бота:\n{str(context.error)[:500]}"
            )
        except:
            pass

# ============== РЕГИСТРАЦИЯ СОТРУДНИКОВ ==============

async def register_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    args = context.args

    if not args:
        await update.message.reply_text(
            "ℹ️ Использование: /register ФИО\n\n"
            "Пример: /register Иванов И.И.\n\n"
            "Если вы уже зарегистрированы, эта команда обновит ваше ФИО."
        )
        return

    full_name = " ".join(args).strip()

    if len(full_name) < 3:
        await update.message.reply_text("❌ Слишком короткое ФИО.")
        return

    existing = None
    if USE_POSTGRES and db_manager:
        existing = db_manager.get_employee_by_id(user_id)
    else:
        for name, uid in EMPLOYEES.items():
            if uid == user_id:
                existing = {"full_name": name, "telegram_id": uid}
                break

    if existing:
        if existing["full_name"] != full_name:
            if USE_POSTGRES and db_manager:
                conn = db_manager.get_connection()
                cur = conn.cursor()
                cur.execute("DELETE FROM employees WHERE telegram_id = %s", (user_id,))
                conn.commit()
                cur.close()
                conn.close()
                if db_manager.register_employee(full_name, user_id):
                    for name in list(EMPLOYEES.keys()):
                        if EMPLOYEES[name] == user_id:
                            del EMPLOYEES[name]
                    EMPLOYEES[full_name] = user_id
                    await update.message.reply_text(
                        f"✅ Ваше ФИО обновлено!\n\n"
                        f"📌 Новое ФИО: {full_name}\n"
                        f"🆔 Ваш Telegram ID: {user_id}"
                    )
                    print(f"✅ Обновлено ФИО сотрудника: {full_name} (ID: {user_id})")
                else:
                    await update.message.reply_text("❌ Ошибка обновления ФИО.")
            else:
                for name in list(EMPLOYEES.keys()):
                    if EMPLOYEES[name] == user_id:
                        del EMPLOYEES[name]
                EMPLOYEES[full_name] = user_id
                await update.message.reply_text(
                    f"✅ Ваше ФИО обновлено (локально)!\n"
                    f"📌 Новое ФИО: {full_name}"
                )
        else:
            await update.message.reply_text(
                f"ℹ️ Вы уже зарегистрированы как {existing['full_name']}.\n"
                f"Если хотите изменить ФИО, используйте /register Новое ФИО"
            )
        return

    if USE_POSTGRES and db_manager:
        existing_by_name = db_manager.get_employee_by_name(full_name)
        if existing_by_name:
            await update.message.reply_text(
                f"❌ ФИО '{full_name}' уже зарегистрировано другим пользователем."
            )
            return

        if db_manager.register_employee(full_name, user_id):
            EMPLOYEES[full_name] = user_id
            await update.message.reply_text(
                f"✅ Вы успешно зарегистрированы!\n\n"
                f"📌 ФИО: {full_name}\n"
                f"🆔 Ваш Telegram ID: {user_id}"
            )
            print(f"✅ Новый сотрудник зарегистрирован в БД: {full_name} (ID: {user_id})")
        else:
            await update.message.reply_text("❌ Ошибка регистрации.")
    else:
        if full_name in EMPLOYEES:
            await update.message.reply_text(f"❌ ФИО '{full_name}' уже занято.")
            return
        EMPLOYEES[full_name] = user_id
        await update.message.reply_text(
            f"✅ Вы зарегистрированы (локально).\n"
            f"📌 ФИО: {full_name}"
        )

# ============== НАСТРОЙКА КОМАНД ДЛЯ ИНТЕРФЕЙСА ==============

async def setup_bot_commands(application: Application):
    commands = [
        ("start", "Запустить бота и показать команды"),
        ("upload", "Загрузить DOCX или PDF"),
        ("list", "Список активных поручений"),
        ("today", "Поручения с дедлайном сегодня"),
        ("week", "Поручения на неделю"),
        ("overdue", "Просроченные поручения"),
        ("completed", "Выполненные поручения"),
        ("responsible", "Фильтр по ответственному"),
        ("deadline", "Фильтр по дате или диапазону"),
        ("menu", "Кнопочное меню"),
        ("search", "Поиск по текстам"),
        ("stats", "Статистика"),
        ("delete", "Удалить текст по ID"),
        ("edit", "Изменить поле (deadline, responsible...)"),
        ("clear", "Удалить все свои поручения"),
        ("export_csv", "Выгрузить поручения в CSV"),
        ("done", "Отметить поручение выполненным"),
        ("register", "Зарегистрироваться для уведомлений"),
        ("remind", "Напомнить о поручении"),
        ("help", "Помощь"),
        ("app", "Открыть приложение"),
    ]
    await application.bot.set_my_commands(commands)

async def app_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Открывает мини-приложение"""
    ngrok_url = "https://entity-underpass-agreed.ngrok-free.dev"  # ← твоя ссылка
    web_app_url = f"{ngrok_url}/tg/app/"
    keyboard = [
        [InlineKeyboardButton("📱 Открыть приложение", web_app={"url": web_app_url})]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(
        "📱 Нажми на кнопку, чтобы открыть мини-приложение:",
        reply_markup=reply_markup
    )
# ============== ЗАПУСК ==============

def main():
    if not TELEGRAM_BOT_TOKEN:
        print("❌ TELEGRAM_BOT_TOKEN не найден в .env")
        return

    app = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    app.post_init = setup_bot_commands

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("help", help_command))
    app.add_handler(CommandHandler("register", register_command))
    app.add_handler(CommandHandler("list", list_assignments))
    app.add_handler(CommandHandler("delete", delete_command))
    app.add_handler(CommandHandler("edit", edit_command))
    app.add_handler(CommandHandler("search", search_command))
    app.add_handler(CommandHandler("stats", stats_command))
    app.add_handler(CommandHandler("today", today_assignments))
    app.add_handler(CommandHandler("week", week_assignments))
    app.add_handler(CommandHandler("overdue", overdue_assignments))
    app.add_handler(CommandHandler("export", export_command))
    app.add_handler(CommandHandler("export_csv", export_csv_command))
    app.add_handler(CommandHandler("clear", clear_command))
    app.add_handler(CommandHandler("done", done_command))
    app.add_handler(CommandHandler("remind", remind_command))
    app.add_handler(CommandHandler("completed", completed_assignments))
    app.add_handler(CommandHandler("responsible", responsible_command))
    app.add_handler(CommandHandler("deadline", deadline_command))
    app.add_handler(CommandHandler("menu", menu_command))
    app.add_handler(CommandHandler("app", app_command))

    app.add_handler(CallbackQueryHandler(complete_callback, pattern="^complete_"))
    app.add_handler(CallbackQueryHandler(responsible_callback, pattern="^resp_"))
    app.add_handler(CallbackQueryHandler(menu_callback_handler, pattern="^menu_"))
    app.add_handler(CallbackQueryHandler(duplicate_callback_handler, pattern="^dup_"))
    app.add_handler(CallbackQueryHandler(review_callback, pattern="^review_"))
    app.add_handler(CallbackQueryHandler(list_navigation, pattern="^list_"))

    # ===== ОБЪЕДИНЁННЫЙ ОБРАБОТЧИК ТЕКСТОВЫХ СООБЩЕНИЙ =====
    async def universal_text_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
        if context.user_data.get("editing_index") is not None:
            await review_edit_handler(update, context)
        elif context.user_data.get("pending_complete") is not None:
            await completion_report_handler(update, context)

    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, universal_text_handler))

    # Файлы и фото – после текста
    app.add_handler(MessageHandler(filters.PHOTO, photo_or_doc_handler))
    app.add_handler(MessageHandler(filters.Document.ALL, photo_or_doc_handler))

    async def upload_command_wrapper(update: Update, context: ContextTypes.DEFAULT_TYPE):
        user_id = update.effective_user.id
        if not is_user_registered(user_id):
            await update.message.reply_text("⚠️ Зарегистрируйтесь: /register ФИО")
            return
        if not is_admin(user_id):
            await update.message.reply_text("⛔ Только начальники могут загружать документы.")
            return
        await upload_command(update, context)

    app.add_handler(CommandHandler("upload", upload_command_wrapper))

    app.add_error_handler(error_handler)

    job_queue = app.job_queue
    if job_queue:
        job_queue.run_daily(
            daily_reminder,
            time=datetime_time(hour=8, minute=30),
            name="daily_reminder"
        )
        print("⏰ Напоминания запланированы на 8:30 каждый день")
    else:
        print("⚠️ JobQueue не установлен. Установи: pip install 'python-telegram-bot[job-queue]'")

    print("🚀 Бот запущен!")
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == "__main__":
    main()